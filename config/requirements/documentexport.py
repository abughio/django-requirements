from docx import *
from docx.shared import Inches

from django.conf import settings
from datetime import date
from .models import Artifact, Project
from django.shortcuts import get_object_or_404
from ast import literal_eval


from .models import *




def DocumentToExport(projectid):
    document = Document()

    # ---- Cover Letter ----
    document.add_picture((r'%s/static/images/my-header.png' % (settings.BASE_DIR)), width=Inches(4))
    document.add_paragraph()
    document.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    #document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()
    project = get_object_or_404(Project, id=projectid)
    artifacts = Artifact.objects.filter(project=project)
    artifacttypes = ArtifactType.objects.all()

    for artifacttype in artifacttypes:
        document.add_heading(artifacttype.name,level=1)
        for artifact in artifacts:
            if artifact.type == artifacttype:
                document.add_heading(artifact.name,level=2)
                fieldvalues = artifact.fieldvalues
                if fieldvalues:
                    table = document.add_table(rows=1, cols=2)
                    table_header = table.rows[0].cells
                    table_header[0].text = 'Label'
                    table_header[1].text = 'Description'
                    for fieldvalue in fieldvalues.all():
                        table_row = table.add_row().cells
                        table_row[0].text = fieldvalue.artifactField.field.label
                        if fieldvalue.artifactField.field.widgetType == 'Radio' or fieldvalue.artifactField.field.widgetType == 'Drop Down':
                            value = FieldValues.objects.filter(id=fieldvalue.value).first()
                            table_row[1].text = value.value
                        elif fieldvalue.artifactField.field.widgetType == 'Checkbox':
                            values_str = ''
                            for i,v in enumerate(literal_eval(fieldvalue.value)):
                                if i == 0:
                                    values_str = FieldValues.objects.filter(id=v).first().value
                                else:
                                    values_str =  values_str+" , "+ FieldValues.objects.filter(id=v).first().value
                            table_row[1].text = values_str

                        else:
                            table_row[1].text = fieldvalue.value


    return document